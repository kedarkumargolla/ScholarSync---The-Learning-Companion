import os
import shutil
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, UnstructuredPowerPointLoader, UnstructuredExcelLoader, CSVLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings, ChatOllama
from langchain_chroma import Chroma
from langchain_core.prompts import PromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser

# Import our new modules
from image_processing import generate_image_description, create_image_document
from table_processing import create_documents_from_dataframe

# Constants
PERSIST_DIRECTORY = "./chroma_db"
COLLECTION_NAME = "local_knowledge_base"
MODEL_NAME = "llama3.2:1b" # Text-only model for non-vision tasks
MULTIMODAL_MODEL = "llava:7b" # Vision-capable model for images
EMBEDDING_MODEL = "nomic-embed-text" # Good for RAG

def get_llm(model_name=MODEL_NAME):
    """Get text-only LLM"""
    return ChatOllama(model=model_name)

def get_multimodal_llm(model_name=MULTIMODAL_MODEL):
    """Get vision-capable LLM for image processing"""
    return ChatOllama(model=model_name)

def get_embeddings(model_name=EMBEDDING_MODEL):
    return OllamaEmbeddings(model=model_name)

def initialize_vectorstore():
    embeddings = get_embeddings()
    if not os.path.exists(PERSIST_DIRECTORY):
        os.makedirs(PERSIST_DIRECTORY)
    
    vectorstore = Chroma(
        persist_directory=PERSIST_DIRECTORY,
        embedding_function=embeddings,
        collection_name=COLLECTION_NAME
    )
    return vectorstore

def clean_pdf_content(docs):
    """Remove lines containing 'BPGP 2024-26 Batch' from PDF documents"""
    cleaned_docs = []
    for doc in docs:
        lines = doc.page_content.split('\n')
        filtered_lines = [line for line in lines if 'BPGP 2024-26 Batch' not in line]
        doc.page_content = '\n'.join(filtered_lines)
        cleaned_docs.append(doc)
    return cleaned_docs

def process_pdf(file_path):
    loader = PyPDFLoader(file_path)
    docs = loader.load()
    return clean_pdf_content(docs)

def process_docx(file_path):
    """Process DOCX and remove headers/footers"""
    try:
        from docx import Document
        import tempfile
        
        # Load the document
        doc = Document(file_path)
        
        # Remove headers
        for section in doc.sections:
            section.header.is_linked_to_previous = False
            for paragraph in section.header.paragraphs:
                paragraph.clear()
        
        # Remove footers
        for section in doc.sections:
            section.footer.is_linked_to_previous = False
            for paragraph in section.footer.paragraphs:
                paragraph.clear()
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        # Load the cleaned document
        loader = Docx2txtLoader(tmp_path)
        docs = loader.load()
        
        # Clean up temp file
        os.remove(tmp_path)
        
        return docs
    except ImportError:
        # Fallback to regular loading if python-docx not available
        loader = Docx2txtLoader(file_path)
        return loader.load()

def process_ppt(file_path):
    loader = UnstructuredPowerPointLoader(file_path)
    return loader.load()

def process_excel(file_path):
    """Process Excel files with enhanced pandas-based chunking"""
    try:
        # Use our new table processing module for better handling
        return create_documents_from_dataframe(file_path, chunk_size=50)
    except Exception as e:
        # Fallback to basic loading if enhanced processing fails
        print(f"Warning: Enhanced XLSX processing failed, using fallback: {e}")
        loader = UnstructuredExcelLoader(file_path)
        return loader.load()

def process_csv(file_path):
    """Process CSV files with enhanced pandas-based chunking"""
    try:
        # Use our new table processing module for better handling
        return create_documents_from_dataframe(file_path, chunk_size=50)
    except Exception as e:
        # Fallback to basic loading if enhanced processing fails
        print(f"Warning: Enhanced CSV processing failed, using fallback: {e}")
        loader = CSVLoader(file_path)
        return loader.load()

def process_image(file_path):
    """Process image files using multimodal LLM"""
    try:
        # Get vision-capable model
        llm = get_multimodal_llm()
        
        # Generate description
        description = generate_image_description(file_path, llm)
        
        # Create document with description
        return [create_image_document(file_path, description)]
    except Exception as e:
        # If description fails, create basic document without it
        print(f"Warning: Image description failed, creating basic document: {e}")
        return [create_image_document(file_path, description=None)]

def load_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    
    handlers = {
        ".pdf": process_pdf,
        ".docx": process_docx,
        ".pptx": process_ppt,
        ".ppt": process_ppt,
        ".xlsx": process_excel,
        ".csv": process_csv,
        # Image formats
        ".jpg": process_image,
        ".jpeg": process_image,
        ".png": process_image,
        ".gif": process_image,
        ".webp": process_image
    }
    
    if ext in handlers:
        return handlers[ext](file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def load_and_split_documents(file_paths, progress_callback=None):
    all_docs = []
    failed_files = []
    ignored_files = []
    
    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".pptx", ".ppt", ".xlsx", ".csv", ".jpg", ".jpeg", ".png", ".gif", ".webp"}
    total_files = len(file_paths)
    
    for i, file_path in enumerate(file_paths):
        filename = os.path.basename(file_path)
        
        if progress_callback:
            progress_callback(i, total_files, filename)
            
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in ALLOWED_EXTENSIONS:
            ignored_files.append(filename)
            continue
            
        try:
            docs = load_file(file_path)
            all_docs.extend(docs)
        except Exception as e:
            failed_files.append(f"{filename}: {str(e)}")
            continue
            
    if not all_docs:
        return {"status": "error", "message": "No valid documents loaded.", "failed": failed_files, "ignored": ignored_files, "splits": []}

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    splits = text_splitter.split_documents(all_docs)
    
    return {
        "status": "success",
        "splits": splits,
        "failed": failed_files,
        "ignored": ignored_files
    }

def index_documents(splits):
    if not splits:
        return
        
    vectorstore = initialize_vectorstore()
    vectorstore.add_documents(documents=splits)

def ingest_files(file_paths):
    # Backward compatibility wrapper
    result = load_and_split_documents(file_paths)
    if result["status"] == "success":
        index_documents(result["splits"])
        
        success_msg = f"Successfully ingested {len(file_paths) - len(result['failed']) - len(result['ignored'])} files ({len(result['splits'])} chunks)."
        if result['ignored']:
            success_msg += f" Ignored: {', '.join(result['ignored'])}."
            
        return {
            "status": "success", 
            "message": success_msg,
            "failed": result['failed']
        }
    return result

def query_rag(question):
    vectorstore = initialize_vectorstore()
    retriever = vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 4})
    
    llm = get_llm()
    
    template = """Your a MBA Knowledge Base. Use the following pieces of context to answer the question at the end to the students. 
    If you don't know the answer, just say that you don't know, don't try to make up an answer. 
    Use three sentences maximum and keep the answer concise.
    
    Context: {context}
    
    Question: {question}
    
    Helpful Answer:"""
    
    prompt = PromptTemplate.from_template(template)
    
    # Helper function to format documents
    def format_docs(docs):
        return "\n\n".join(doc.page_content for doc in docs)
    
    # Get source documents first
    source_docs = retriever.invoke(question)
    
    # Create RAG chain using LCEL (LangChain Expression Language)
    rag_chain = (
        {"context": lambda x: format_docs(source_docs), "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )
    
    # Get answer
    answer = rag_chain.invoke(question)
    
    return answer, source_docs

def clear_database():
    if os.path.exists(PERSIST_DIRECTORY):
        shutil.rmtree(PERSIST_DIRECTORY)
        return "Database cleared."
    return "Database already empty."
