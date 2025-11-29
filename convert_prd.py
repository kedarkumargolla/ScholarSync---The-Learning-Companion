#!/usr/bin/env python
"""
Convert PRD.md to DOCX format using python-docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def parse_markdown_to_docx(md_file, docx_file):
    """Parse markdown and create DOCX with formatting"""
    
    # Create document
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    print(f"Processing {len(lines)} lines...")
    
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []
    
    for i, line in enumerate(lines):
        line = line.rstrip()
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_block_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                p_format = p.paragraph_format
                p_format.left_indent = Inches(0.5)
                
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(64, 64, 64)
                
                code_block_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            continue
        
        # Handle tables
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            continue
        elif in_table and table_lines:
            # End of table, process it
            process_table(doc, table_lines)
            table_lines = []
            in_table = False
        
        # Handle headings
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
        elif line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
        elif line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
        elif line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_heading(text, level=4)
        elif line.startswith('##### '):
            text = line[6:].strip()
            p = doc.add_heading(text, level=5)
        
        # Handle horizontal rules
        elif line.startswith('---') or line.startswith('___'):
            p = doc.add_paragraph()
            p.add_run('─' * 80)
            p_format = p.paragraph_format
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(6)
        
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            text = clean_markdown(text)
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line).strip()
            text = clean_markdown(text)
            p = doc.add_paragraph(text, style='List Number')
        
        # Handle blockquotes
        elif line.startswith('> '):
            text = line[2:].strip()
            text = clean_markdown(text)
            p = doc.add_paragraph(text)
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Handle regular paragraphs
        elif line.strip():
            text = clean_markdown(line.strip())
            p = doc.add_paragraph(text)
        
        # Empty line
        else:
            continue
        
        # Progress indicator
        if i % 100 == 0 and i > 0:
            print(f"Processed {i}/{len(lines)} lines...")
    
    # Handle remaining table if any
    if in_table and table_lines:
        process_table(doc, table_lines)
    
    # Save document
    print(f"Saving {docx_file}...")
    doc.save(docx_file)
    print(f"✅ Successfully created {docx_file}")
    
    return True

def clean_markdown(text):
    """Remove markdown formatting from text"""
    # Bold
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Italic
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # Code
    text = re.sub(r'`(.*?)`', r'\1', text)
    # Links
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    return text

def process_table(doc, table_lines):
    """Process markdown table and add to document"""
    if len(table_lines) < 2:
        return
    
    # Parse header
    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    
    # Skip separator line (line with dashes)
    data_lines = [line for line in table_lines[2:] if not line.strip().startswith('|---')]
    
    if not data_lines:
        return
    
    # Create table
    table = doc.add_table(rows=len(data_lines) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = clean_markdown(header)
        # Bold header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add data
    for row_idx, line in enumerate(data_lines):
        cells_data = [cell.strip() for cell in line.split('|')[1:-1]]
        for col_idx, cell_data in enumerate(cells_data):
            if col_idx < len(headers):
                table.rows[row_idx + 1].cells[col_idx].text = clean_markdown(cell_data)

if __name__ == "__main__":
    input_file = "PRD.md"
    output_file = "PRD.docx"
    
    if not os.path.exists(input_file):
        print(f"❌ Error: {input_file} not found!")
        exit(1)
    
    print(f"Converting {input_file} to {output_file}...")
    
    try:
        success = parse_markdown_to_docx(input_file, output_file)
        
        if success and os.path.exists(output_file):
            size_kb = os.path.getsize(output_file) / 1024
            print(f"\n✨ Conversion complete!")
            print(f"📄 Created: {output_file} ({size_kb:.1f} KB)")
            print(f"\n💡 To create PDF:")
            print(f"   1. Open {output_file} in Microsoft Word")
            print(f"   2. File → Save As → PDF")
            print(f"   Or install MiKTeX for automated PDF generation")
        else:
            print(f"❌ Conversion failed!")
            exit(1)
            
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        exit(1)
